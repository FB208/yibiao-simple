import streamlit as st
import json
import threading
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, List, Tuple
import io
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def render_content_edit_page() -> Dict:
    """
    渲染正文编辑页面，支持AI生成标书正文
    
    Returns:
        包含页面状态数据的字典
    """
    
    # 添加返回顶部按钮 - 使用简单的HTML实现
    st.markdown("""
    <style>
        .back-to-top {
            position: fixed;
            bottom: 20px;
            right: 20px;
            background-color: #ff6b6b;
            color: white;
            border: none;
            border-radius: 50%;
            width: 50px;
            height: 50px;
            font-size: 18px;
            cursor: pointer;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
            z-index: 999;
            text-align: center;
            line-height: 50px;
            text-decoration: none;
            display: block;
        }
        .back-to-top:hover {
            background-color: #ff5252;
            text-decoration: none;
            color: white;
        }
    </style>
    <a href="#top" class="back-to-top" title="返回顶部">↑</a>
    """, unsafe_allow_html=True)
    
    # 添加顶部锚点
    st.markdown('<div id="top"></div>', unsafe_allow_html=True)
    
    st.header("📝 正文编辑")
    
    # 显示传入的目录数据（折叠形式，方便调试）
    st.session_state.outline_data = {"outline":[{"id":"1","title":"项目总体概述及运维方案","description":"针对本项目有深刻认识，总体概述清晰、合理，运维方案以及措施先进、成熟完全满足规范要求","children":[{"id":"1.1","title":"项目概述","description":"介绍项目背景、目标、范围、规模与预算等基本情况","children":[{"id":"1.1.1","title":"项目背景与意义","description":"阐述天津港保税区消防救援支队消防物联网远程监控系统运维服务项目的背景和重要意义"},{"id":"1.1.2","title":"项目目标与范围","description":"明确项目运维目标和具体服务范围，包括监控中心、IDC机房、通讯链路、平台软硬件及前端设备"},{"id":"1.1.3","title":"项目规模与预算","description":"详细说明项目规模和预算情况，包括各项资源配置和费用构成"}]},{"id":"1.2","title":"运维方案总体设计","description":"阐述运维服务的总体设计理念和架构","children":[{"id":"1.2.1","title":"运维服务理念与原则","description":"阐述运维服务的核心理念和基本原则"},{"id":"1.2.2","title":"运维服务体系架构","description":"设计完整的运维服务体系架构，包括组织架构、流程架构和技术架构"},{"id":"1.2.3","title":"运维服务流程设计","description":"设计标准化的运维服务流程，确保服务质量和效率"}]},{"id":"1.3","title":"运维服务内容","description":"详细说明各项运维服务的具体内容和实施方式","children":[{"id":"1.3.1","title":"监控中心与IDC机房运维","description":"阐述监控中心场地和IDC机房机柜的运维服务内容"},{"id":"1.3.2","title":"通讯链路运维","description":"说明固定IP专线、物联网卡和视频专线的运维服务内容"},{"id":"1.3.3","title":"平台软硬件运维","description":"阐述系统平台软硬件的日常维护和优化服务内容"},{"id":"1.3.4","title":"前端设备电池更换","description":"详细说明用传蓄电池和传感器电池的更换方案和实施计划"}]}]},{"id":"2","title":"本项目投入人员团队的评价","description":"展示项目实施人员配置、技术实力和同类项目实施或服务经验","children":[{"id":"2.1","title":"人员团队总体构成","description":"介绍人员团队的整体构成情况","children":[{"id":"2.1.1","title":"团队规模与配置","description":"详细说明团队总人数、专业配置和人员结构"},{"id":"2.1.2","title":"团队专业结构","description":"阐述团队成员的专业背景和技术特长分布"},{"id":"2.1.3","title":"团队资质认证","description":"列举团队成员持有的相关资质证书和专业认证"}]},{"id":"2.2","title":"核心技术人员介绍","description":"详细介绍项目核心技术人员的情况","children":[{"id":"2.2.1","title":"项目负责人资质与经验","description":"介绍项目负责人的专业资质、从业经验和项目管理能力"},{"id":"2.2.2","title":"技术骨干资质与经验","description":"介绍技术骨干的专业背景、技术能力和项目经验"},{"id":"2.2.3","title":"消防专业人员资质与经验","description":"介绍持有消防员证书等专业资质的人员情况"}]},{"id":"2.3","title":"团队业绩与经验","description":"展示团队过往的项目业绩和经验积累","children":[{"id":"2.3.1","title":"类似项目实施经验","description":"列举团队参与的类似消防物联网系统项目经验"},{"id":"2.3.2","title":"消防物联网系统运维经验","description":"详述团队在消防物联网系统运维方面的专业经验"},{"id":"2.3.3","title":"技术创新与成果","description":"展示团队在相关领域的技术创新和成果"}]}]},{"id":"3","title":"人员组织方案及分工职责","description":"详细说明人员配备方案、组织安排和职责分工","children":[{"id":"3.1","title":"组织架构设计","description":"设计合理的运维服务组织架构","children":[{"id":"3.1.1","title":"运维服务组织架构","description":"设计完整的运维服务组织架构图和说明"},{"id":"3.1.2","title":"管理层级与汇报关系","description":"明确各管理层级和汇报关系，确保信息畅通"},{"id":"3.1.3","title":"协同工作机制","description":"建立高效的团队协同工作机制，提高工作效率"}]},{"id":"3.2","title":"岗位职责划分","description":"明确各岗位的具体职责和工作内容","children":[{"id":"3.2.1","title":"项目经理职责","description":"详细说明项目经理的职责范围和工作要求"},{"id":"3.2.2","title":"技术负责人职责","description":"明确技术负责人的职责范围和技术管理要求"},{"id":"3.2.3","title":"运维人员职责","description":"划分各类运维人员的具体职责和工作内容"}]},{"id":"3.3","title":"人员配置方案","description":"详细说明各类人员的配置方案和工作安排","children":[{"id":"3.3.1","title":"监控中心值守人员配置","description":"说明8人7X24小时值守人员的配置方案和工作安排"},{"id":"3.3.2","title":"前端设备运维人员配置","description":"说明3名硬件技术人员的配置方案和工作范围"},{"id":"3.3.3","title":"系统软件运维人员配置","description":"说明不少于2名软件技术人员的配置方案和工作内容"}]}]},{"id":"4","title":"前端硬件设备运维方案","description":"全面阐述消防物联网前端硬件设备的运维、维保和维修方案","children":[{"id":"4.1","title":"前端设备概况","description":"介绍前端设备的基本情况和运维难点","children":[{"id":"4.1.1","title":"设备类型与分布","description":"详细列出前端设备的类型、数量和分布情况"},{"id":"4.1.2","title":"设备运行状态分析","description":"分析前端设备的当前运行状态和潜在问题"},{"id":"4.1.3","title":"设备运维难点分析","description":"分析前端设备运维过程中的难点和挑战"}]},{"id":"4.2","title":"设备维保方案","description":"制定前端设备的日常维保方案","children":[{"id":"4.2.1","title":"日常巡检计划","description":"制定详细的日常巡检计划，包括巡检频率、内容和标准"},{"id":"4.2.2","title":"定期维护保养","description":"制定设备定期维护保养计划，确保设备长期稳定运行"},{"id":"4.2.3","title":"设备性能优化","description":"提出设备性能优化方案，提高设备运行效率"}]},{"id":"4.3","title":"设备维修方案","description":"制定前端设备故障维修方案","children":[{"id":"4.3.1","title":"故障诊断流程","description":"建立标准化的故障诊断流程，快速定位问题"},{"id":"4.3.2","title":"维修响应机制","description":"制定高效的维修响应机制，确保及时处理故障"},{"id":"4.3.3","title":"备品备件管理","description":"建立备品备件管理制度，保障维修需求"}]},{"id":"4.4","title":"电池更换方案","description":"详细说明前端设备电池更换方案","children":[{"id":"4.4.1","title":"用传蓄电池更换方案","description":"制定260块用传蓄电池的更换方案和实施计划"},{"id":"4.4.2","title":"传感器电池更换方案","description":"制定1820块传感器电池的更换方案和实施计划"},{"id":"4.4.3","title":"电池更换质量控制","description":"建立电池更换质量控制体系，确保更换质量"}]}]},{"id":"5","title":"软硬件故障应急处理方案","description":"制定全面、合理、可行的软硬件故障应急处理方案","children":[{"id":"5.1","title":"应急处理总体架构","description":"设计应急处理的总体架构和机制","children":[{"id":"5.1.1","title":"应急响应组织","description":"建立应急响应组织架构，明确各岗位职责"},{"id":"5.1.2","title":"应急响应流程","description":"制定标准化的应急响应流程，确保快速响应"},{"id":"5.1.3","title":"应急资源保障","description":"配置必要的应急资源，确保应急处理能力"}]},{"id":"5.2","title":"故障分级与响应机制","description":"建立故障分级标准和相应的响应机制","children":[{"id":"5.2.1","title":"故障等级划分","description":"制定故障等级划分标准，明确各级故障的定义"},{"id":"5.2.2","title":"响应时间要求","description":"规定不同等级故障的响应时间和处理时限"},{"id":"5.2.3","title":"升级处理机制","description":"建立故障升级处理机制，确保重大故障得到及时处理"}]},{"id":"5.3","title":"典型故障处理预案","description":"制定典型故障的处理预案","children":[{"id":"5.3.1","title":"系统软件故障处理预案","description":"制定系统软件故障的应急处理预案"},{"id":"5.3.2","title":"网络通信故障处理预案","description":"制定网络通信故障的应急处理预案"},{"id":"5.3.3","title":"前端设备故障处理预案","description":"制定前端设备故障的应急处理预案"}]},{"id":"5.4","title":"应急演练与评估","description":"建立应急演练和评估机制","children":[{"id":"5.4.1","title":"应急演练计划","description":"制定定期应急演练计划，提高应急处理能力"},{"id":"5.4.2","title":"演练效果评估","description":"建立演练效果评估机制，持续改进应急处理能力"},{"id":"5.4.3","title":"预案持续优化","description":"根据演练结果和实际情况，持续优化应急预案"}]}]},{"id":"6","title":"人员稳定性保障措施","description":"制定完整、切实可行的人员稳定性保障措施","children":[{"id":"6.1","title":"人员管理制度","description":"建立完善的人员管理制度","children":[{"id":"6.1.1","title":"人员招聘与选拔","description":"制定科学的人员招聘与选拔机制，确保人员质量"},{"id":"6.1.2","title":"薪酬福利体系","description":"建立有竞争力的薪酬福利体系，提高员工满意度"},{"id":"6.1.3","title":"职业发展通道","description":"设计清晰的职业发展通道，增强员工归属感"}]},{"id":"6.2","title":"人员激励措施","description":"制定有效的人员激励措施","children":[{"id":"6.2.1","title":"绩效考核机制","description":"建立科学的绩效考核机制，激发员工工作积极性"},{"id":"6.2.2","title":"奖惩制度","description":"制定合理的奖惩制度，引导员工行为"},{"id":"6.2.3","title":"团队建设活动","description":"组织丰富的团队建设活动，增强团队凝聚力"}]},{"id":"6.3","title":"人员替换机制","description":"建立规范的人员替换机制","children":[{"id":"6.3.1","title":"人员替换流程","description":"制定规范的人员替换流程，确保工作连续性"},{"id":"6.3.2","title":"知识转移机制","description":"建立有效的知识转移机制，保障经验传承"},{"id":"6.3.3","title":"工作交接标准","description":"制定标准化的工作交接规范，确保无缝衔接"}]}]},{"id":"7","title":"定期预防性检查方案","description":"制定细致、全面、可操作性强的定期预防性检查方案","children":[{"id":"7.1","title":"预防性检查体系","description":"建立完整的预防性检查体系","children":[{"id":"7.1.1","title":"检查周期与频率","description":"制定合理的检查周期和频率，确保及时发现问题"},{"id":"7.1.2","title":"检查内容与标准","description":"明确检查内容和标准，确保检查质量"},{"id":"7.1.3","title":"检查方法与工具","description":"采用科学的检查方法和工具，提高检查效率"}]},{"id":"7.2","title":"系统软件检查方案","description":"制定系统软件的定期检查方案","children":[{"id":"7.2.1","title":"平台系统健康检查","description":"制定平台系统健康检查方案，确保系统稳定运行"},{"id":"7.2.2","title":"数据库性能检查","description":"制定数据库性能检查方案，优化数据处理效率"},{"id":"7.2.3","title":"安全漏洞检查","description":"制定安全漏洞检查方案，保障系统安全"}]},{"id":"7.3","title":"硬件设备检查方案","description":"制定硬件设备的定期检查方案","children":[{"id":"7.3.1","title":"监控中心设备检查","description":"制定监控中心设备检查方案，确保设备正常运行"},{"id":"7.3.2","title":"通信链路检查","description":"制定通信链路检查方案，保障数据传输畅通"},{"id":"7.3.3","title":"前端设备检查","description":"制定前端设备检查方案，确保设备状态良好"}]},{"id":"7.4","title":"检查结果处理与改进","description":"建立检查结果处理和持续改进机制","children":[{"id":"7.4.1","title":"问题分级与处理","description":"建立问题分级处理机制，确保问题得到及时解决"},{"id":"7.4.2","title":"改进措施跟踪","description":"建立改进措施跟踪机制，确保改进措施落实到位"},{"id":"7.4.3","title":"检查报告与总结","description":"制定检查报告和总结制度，为决策提供依据"}]}]}]}
    if st.session_state.get('outline_data'):
        with st.expander("📋 查看传入的目录数据", expanded=False):
            st.json(st.session_state.outline_data)
    else:
        st.warning("⚠️ 未检测到目录数据，请先完成目录编辑步骤")
        return {}
    
    
    # 创建两列布局
    col1, col2 = st.columns([3, 1])
    
    with col2:
        st.markdown("**操作面板**")
        
        # 生成正文按钮
        generate_button = st.button("🤖 生成正文", use_container_width=True, type="primary")
        
        # 导出按钮（随时可以点击，方便调试）
        export_button = st.button("📤 导出Word文档", use_container_width=True)
        
        # 调试信息（可选，帮助排查问题）
        if st.checkbox("🔍 显示调试信息", value=False):
            generated_content = st.session_state.get('generated_content', {})
            st.write(f"generated_content 存在: {generated_content is not None}")
            st.write(f"generated_content 长度: {len(generated_content) if generated_content else 0}")
            if generated_content:
                st.write("generated_content 键:", list(generated_content.keys())[:3])  # 显示前3个键
        
        if export_button:
            _export_document()
    
    with col1:
        # 内容显示区域
        content_container = st.container()
        
        if generate_button:
            _generate_content(content_container)
        else:
            # 显示已生成的内容（如果有）
            if st.session_state.get('generated_content'):
                _display_generated_content(content_container)
            else:
                with content_container:
                    st.info("点击\"生成正文\"按钮开始AI生成标书内容")
    
    return {
        'outline_data': st.session_state.get('outline_data'),
        'generated_content': st.session_state.get('generated_content'),
        'content_generated': st.session_state.get('content_generated', False)
    }


def _generate_content(container):
    """
    生成正文内容的主要逻辑
    """
    try:
        from services.openai_servce import get_openai_service
        
        with container:
            st.markdown("### 🔄 正在生成正文内容...")
            
            # 初始化服务
            openai_service = get_openai_service()
            outline_data = st.session_state.get('outline_data', {})
            
            project_overview = st.session_state.get('project_overview', '')
            
            # 创建进度条
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # 初始化生成内容存储
            if 'generated_content' not in st.session_state:
                st.session_state.generated_content = {}
            
            # 计算总的叶子节点数
            total_nodes = _count_leaf_nodes(outline_data)
            
            if total_nodes == 0:
                st.error("未找到可生成内容的叶子节点")
                return
            
            # 使用并发方式处理outline_data，生成叶子节点内容
            _process_outline_concurrent(
                container,
                outline_data,
                openai_service,
                project_overview,
                progress_bar,
                status_text,
                total_nodes
            )
            
            # 生成完成
            progress_bar.progress(1.0)
            status_text.success("✅ 所有章节内容生成完成！")
            st.session_state.content_generated = True
            
            # 显示完成摘要
            st.success(f"🎉 成功生成 {total_nodes} 个章节的内容")
            
            # 刷新页面状态，确保导出按钮能立即可用
            st.rerun()
            
    except ImportError:
        st.error("❌ 无法导入OpenAI服务，请检查服务配置")
    except Exception as e:
        st.error(f"❌ 生成过程中发生错误: {str(e)}")


def _display_generated_content(container):
    """
    显示已生成的内容
    """
    with container:
        st.markdown("### 📋 生成内容预览")
        
        generated_content = st.session_state.get('generated_content', {})
        outline_data = st.session_state.get('outline_data', {})
        
        if not generated_content:
            st.info("暂无已生成的内容")
            return
        
        # 按照目录结构重新显示内容，与生成时的显示方式保持一致
        _display_content_by_outline_structure(generated_content, outline_data)


def _display_content_by_outline_structure(generated_content, outline_data):
    """
    按照目录结构显示已生成的内容，与生成时的显示方式保持一致
    
    Args:
        generated_content: 生成的内容字典
        outline_data: 目录数据
    """
    if not outline_data or 'outline' not in outline_data:
        # 如果没有目录结构，就按简单方式显示
        for node_path, content in generated_content.items():
            # 解析路径获取标题
            parts = node_path.split(' ', 1)
            title = parts[1] if len(parts) > 1 else node_path
            st.markdown(f"### {title}")
            st.markdown(content)
            st.markdown("---")
        return
    
    # 按照目录结构递归显示
    for chapter in outline_data['outline']:
        _display_chapter_content_recursive(chapter, generated_content, level=1)


def _display_chapter_content_recursive(chapter, generated_content, level=1):
    """
    递归显示章节内容，保持与生成时的显示格式一致
    
    Args:
        chapter: 当前章节
        generated_content: 生成的内容字典
        level: 当前层级
    """
    chapter_id = chapter.get('id', '')
    chapter_title = chapter.get('title', '未命名章节')
    
    # 显示章节标题
    title_prefix = "#" * min(level + 2, 6) + " "
    st.markdown(title_prefix + chapter_title)
    
    # 检查是否为叶子节点
    is_leaf = not chapter.get('children') or len(chapter['children']) == 0
    
    if is_leaf:
        # 叶子节点，显示生成的内容
        chapter_path = f"{chapter_id} {chapter_title}"
        if chapter_path in generated_content:
            content = generated_content[chapter_path]
            if content and not content.startswith('生成'):  # 排除错误信息
                st.markdown(content)
            else:
                st.error(content)  # 显示错误信息
    else:
        # 非叶子节点，递归处理子节点
        for child in chapter['children']:
            _display_chapter_content_recursive(child, generated_content, level + 1)


def _count_leaf_nodes(data) -> int:
    """
    计算outline_data中叶子节点的总数
    
    Args:
        data: outline数据
        
    Returns:
        叶子节点总数
    """
    if not data or 'outline' not in data:
        return 0
    
    return _count_leaf_nodes_recursive(data['outline'])


def _count_leaf_nodes_recursive(nodes) -> int:
    """
    递归计算叶子节点数量
    """
    count = 0
    for node in nodes:
        # 检查是否为叶子节点
        if not node.get('children') or len(node['children']) == 0:
            count += 1
        else:
            count += _count_leaf_nodes_recursive(node['children'])
    
    return count


def _process_outline_concurrent(container, data, openai_service, project_overview, progress_bar, status_text, total_nodes):
    """
    并发处理outline_data，为叶子节点生成内容
    
    Args:
        container: Streamlit容器
        data: outline数据
        openai_service: OpenAI服务实例
        project_overview: 项目概述
        progress_bar: 进度条组件
        status_text: 状态文本组件
        total_nodes: 总叶子节点数
    """
    with container:
        # 第一步：收集所有叶子节点信息和创建显示容器
        leaf_nodes_info = []
        node_containers = {}
        
        st.markdown("### 📋 生成内容预览")
        
        # 递归收集叶子节点信息并创建显示结构
        _collect_and_display_structure(data, leaf_nodes_info, node_containers, parent_chapters=[], level=1)
        
        if not leaf_nodes_info:
            st.error("未找到叶子节点")
            return
        
        # 第二步：并发生成内容
        _generate_content_concurrent(leaf_nodes_info, node_containers, openai_service, project_overview, progress_bar, status_text)


def _collect_and_display_structure(data, leaf_nodes_info, node_containers, parent_chapters=None, level=1):
    """
    递归收集叶子节点信息并创建显示结构
    
    Args:
        data: 当前处理的数据
        leaf_nodes_info: 叶子节点信息列表
        node_containers: 节点容器字典
        parent_chapters: 父级章节信息
        level: 当前层级
    """
    if parent_chapters is None:
        parent_chapters = []
    
    # 处理outline根节点
    if isinstance(data, dict) and 'outline' in data:
        outline_chapters = data['outline']
        for chapter in outline_chapters:
            _process_chapter_structure_with_siblings(chapter, outline_chapters, leaf_nodes_info, node_containers, parent_chapters, level)
    
    


def _process_chapter_structure_with_siblings(chapter, all_siblings, leaf_nodes_info, node_containers, parent_chapters, level):
    """
    处理单个章节的结构，包含兄弟节点信息
    """
    chapter_id = chapter.get('id', 'unknown')
    chapter_title = chapter.get('title', '未命名章节')
    chapter_path = f"{chapter_id} {chapter_title}"
    
    # 显示当前章节标题
    title_prefix = "#" * min(level + 2, 6) + " "
    st.markdown(title_prefix + chapter_title)
    
    # 检查是否为叶子节点
    is_leaf = not chapter.get('children') or len(chapter['children']) == 0
    
    if is_leaf:
        # 叶子节点：创建容器并收集信息
        content_container = st.empty()
        node_containers[chapter_path] = content_container
        
        # 收集兄弟节点信息（排除自己）
        sibling_chapters = [sib for sib in all_siblings if sib.get('id') != chapter.get('id')]
        
        # 构建节点信息
        node_info = {
            'chapter': chapter,
            'parent_chapters': parent_chapters.copy(),
            'sibling_chapters': sibling_chapters,
            'level': level,
            'path': chapter_path
        }
        leaf_nodes_info.append((chapter_path, node_info))
        
        # 在容器中显示"生成中..."提示
        content_container.info("🔄 等待生成...")
        
    else:
        # 非叶子节点：递归处理子节点
        current_chapter_info = {
            'id': chapter.get('id'),
            'title': chapter.get('title'),
            'description': chapter.get('description', '')
        }
        new_parent_chapters = parent_chapters + [current_chapter_info]
        
        # 处理所有子节点
        children = chapter['children']
        for child in children:
            _process_chapter_structure_with_siblings(child, children, leaf_nodes_info, node_containers, new_parent_chapters, level + 1)


def _generate_content_concurrent(leaf_nodes_info, node_containers, openai_service, project_overview, progress_bar, status_text):
    """
    并发生成内容
    """
    total_nodes = len(leaf_nodes_info)
    
    # 预先提取配置信息，避免在子线程中访问
    api_key = openai_service.api_key
    base_url = openai_service.base_url
    model_name = openai_service.model_name
    
    def generate_single_node_safe(node_data):
        """线程安全的内容生成函数 - 不访问任何Streamlit组件"""
        node_path, node_info = node_data
        try:
            # 在子线程中创建独立的服务实例
            from services.openai_servce import OpenAIService
            thread_service = OpenAIService(
                api_key=api_key,
                base_url=base_url,
                model_name=model_name
            )
            
            # 生成内容
            generated_text = thread_service._generate_chapter_content(
                node_info['chapter'],
                parent_chapters=node_info['parent_chapters'],
                sibling_chapters=node_info['sibling_chapters'],
                project_overview=project_overview
            )
            
            return node_path, generated_text, None
            
        except Exception as e:
            error_msg = f"生成失败: {str(e)}"
            return node_path, None, error_msg
    
    # 使用线程池并发处理，但在主线程中更新UI
    max_workers = 5
    completed_count = 0
    
    # 创建进度显示容器
    progress_container = st.container()
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 提交所有任务
        futures = {executor.submit(generate_single_node_safe, node_data): node_data 
                  for node_data in leaf_nodes_info}
        
        # 处理完成的任务
        for future in concurrent.futures.as_completed(futures):
            node_data = futures[future]
            node_path = node_data[0]
            
            try:
                result_path, content, error = future.result()
                
                # 主线程中更新UI
                container = node_containers[result_path]
                if error:
                    container.error(error)
                else:
                    container.markdown(content)
                
                # 保存到session state
                if 'generated_content' not in st.session_state:
                    st.session_state.generated_content = {}
                st.session_state.generated_content[result_path] = content if not error else error
                
                # 更新进度
                completed_count += 1
                progress = completed_count / total_nodes
                progress_bar.progress(progress)
                status_text.markdown(f"**已完成 {completed_count}/{total_nodes} 个章节**")
                
            except Exception as exc:
                completed_count += 1
                progress = completed_count / total_nodes
                progress_bar.progress(progress)
                
                # 显示错误
                container = node_containers[node_path]
                container.error(f'生成异常: {exc}')
                
                # 保存错误信息
                if 'generated_content' not in st.session_state:
                    st.session_state.generated_content = {}
                st.session_state.generated_content[node_path] = f'生成异常: {exc}'


def _export_document():
    """
    导出生成的文档内容为Word格式
    """
    try:
        generated_content = st.session_state.get('generated_content', {})
        outline_data = st.session_state.get('outline_data', {})
        
        if not generated_content:
            st.warning("⚠️ 暂无已生成的内容，将导出空白文档模板")
            generated_content = {}
        
        if not DOCX_AVAILABLE:
            st.error("❌ 缺少python-docx库，无法导出Word文档。请运行: pip install python-docx")
            return
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体为宋体
        _set_document_font(doc, '宋体')
        
        # 添加标题
        title = doc.add_heading('标书正文', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_font(title, '宋体', 18)
        
        # 添加AI生成说明
        ai_para = doc.add_paragraph('内容由AI生成')
        _set_paragraph_font(ai_para, '宋体', 12)
        
        # 添加生成时间
        time_para = doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        _set_paragraph_font(time_para, '宋体', 12)
        
        doc.add_paragraph('')  # 空行
        
        # 按照目录结构添加内容
        if outline_data and 'outline' in outline_data:
            for chapter in outline_data['outline']:
                _add_chapter_to_word_doc(doc, chapter, generated_content, level=1)
        else:
            # 如果没有目录结构，直接添加所有内容
            for node_path, content in generated_content.items():
                parts = node_path.split(' ', 1)
                title = parts[1] if len(parts) > 1 else node_path
                heading_para = doc.add_heading(title, 2)
                _set_paragraph_font(heading_para, '宋体', 14)
                
                content_para = doc.add_paragraph(content)
                _set_paragraph_font(content_para, '宋体', 12)
        
        # 将文档保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # 创建文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"标书正文_{timestamp}.docx"
        
        # 提供下载
        st.download_button(
            label="💾 下载Word文档",
            data=doc_buffer.getvalue(),
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        # 显示成功信息
        content_count = len(generated_content)
        st.success(f"✅ Word文档已准备完成！包含 {content_count} 个已生成章节。点击上方按钮下载 {filename}")
        
        # 显示内容统计
        if generated_content:
            with st.expander("📊 导出内容统计", expanded=False):
                st.write(f"**包含章节数量:** {content_count}")
                st.write("**包含章节列表:**")
                for i, node_path in enumerate(generated_content.keys(), 1):
                    st.write(f"{i}. {node_path}")
        else:
            st.info("ℹ️ 导出的是空白文档模板，包含标题和时间信息")
            
    except Exception as e:
        st.error(f"❌ 导出过程中发生错误: {str(e)}")


def _set_document_font(doc, font_name):
    """
    设置整个文档的默认字体
    
    Args:
        doc: Word文档对象
        font_name: 字体名称
    """
    # 设置样式的默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(12)
    
    # 为中文设置字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _set_paragraph_font(paragraph, font_name, size=12):
    """
    设置段落的字体
    
    Args:
        paragraph: 段落对象
        font_name: 字体名称
        size: 字体大小
    """
    # 如果段落没有runs，添加一个空的run
    if not paragraph.runs:
        run = paragraph.add_run()
    
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(size)
        # 为中文设置字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _add_chapter_to_word_doc(doc, chapter, generated_content, level=1):
    """
    递归添加章节内容到Word文档
    
    Args:
        doc: Word文档对象
        chapter: 当前章节
        generated_content: 生成的内容字典
        level: 当前层级
    """
    chapter_id = chapter.get('id', '')
    chapter_title = chapter.get('title', '未命名章节')
    chapter_description = chapter.get('description', '')
    
    # 添加章节标题（Word标题级别从1开始，最大到9）
    heading_level = min(level + 1, 9)
    full_title = f"{chapter_id} {chapter_title}"
    heading_para = doc.add_heading(full_title, heading_level)
    # 设置标题字体为宋体
    heading_size = max(16 - level, 12)  # 根据层级设置字体大小
    _set_paragraph_font(heading_para, '宋体', heading_size)
    
    # 检查是否为叶子节点
    is_leaf = not chapter.get('children') or len(chapter['children']) == 0
    
    if is_leaf:
        # 叶子节点，添加生成的内容
        chapter_path = f"{chapter_id} {chapter_title}"
        if chapter_path in generated_content:
            content = generated_content[chapter_path]
            if content and not content.startswith('生成'):  # 排除错误信息
                # 清理markdown格式并添加到Word文档
                cleaned_content = _clean_markdown_for_word(content)
                content_para = doc.add_paragraph(cleaned_content)
                _set_paragraph_font(content_para, '宋体', 12)
            else:
                # 显示错误信息
                error_para = doc.add_paragraph(f"[生成错误: {content}]")
                _set_paragraph_font(error_para, '宋体', 12)
        else:
            # 如果没有生成内容，添加描述或占位符
            if chapter_description:
                placeholder_para = doc.add_paragraph(f"[待完善: {chapter_description}]")
                _set_paragraph_font(placeholder_para, '宋体', 12)
            else:
                placeholder_para = doc.add_paragraph("[待完善]")
                _set_paragraph_font(placeholder_para, '宋体', 12)
    else:
        # 非叶子节点，添加描述（如果有）并递归处理子节点
        if chapter_description:
            desc_para = doc.add_paragraph(chapter_description)
            desc_para.italic = True
            _set_paragraph_font(desc_para, '宋体', 12)
        
        # 递归处理子节点
        for child in chapter['children']:
            _add_chapter_to_word_doc(doc, child, generated_content, level + 1)


def _clean_markdown_for_word(content):
    """
    清理markdown格式，使其适合Word文档
    
    Args:
        content: 原始markdown内容
        
    Returns:
        清理后的文本内容
    """
    if not content:
        return ""
    
    # 简单的markdown清理，去除常见的markdown标记
    import re
    
    # 去除markdown标题标记
    content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)
    
    # 去除markdown粗体和斜体标记
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # 粗体
    content = re.sub(r'\*(.*?)\*', r'\1', content)      # 斜体
    
    # 去除markdown链接标记，保留链接文本
    content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)
    
    # 去除markdown列表标记
    content = re.sub(r'^[-*+]\s+', '', content, flags=re.MULTILINE)
    content = re.sub(r'^\d+\.\s+', '', content, flags=re.MULTILINE)
    
    # 去除markdown代码块标记
    content = re.sub(r'```.*?```', '', content, flags=re.DOTALL)
    content = re.sub(r'`(.*?)`', r'\1', content)
    
    return content.strip()


def _organize_content_by_outline(generated_content, outline_data):
    """
    按照目录结构组织生成的内容
    
    Args:
        generated_content: 生成的内容字典
        outline_data: 目录数据
        
    Returns:
        组织后的内容结构
    """
    organized = {}
    
    # 遍历generated_content，按路径组织
    for path, content in generated_content.items():
        # 解析路径，提取章节ID
        parts = path.split(' ', 1)
        if len(parts) >= 2:
            chapter_id = parts[0]
            chapter_title = parts[1]
            organized[chapter_id] = {
                'title': chapter_title,
                'content': content,
                'path': path
            }
    
    return organized


def _generate_document_content(organized_content, outline_data):
    """
    生成完整的文档内容
    
    Args:
        organized_content: 组织后的内容
        outline_data: 目录数据
        
    Returns:
        完整的文档字符串
    """
    doc_lines = []
    
    # 添加文档标题
    doc_lines.append("# 标书正文")
    doc_lines.append("")
    doc_lines.append(f"*生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}*")
    doc_lines.append("")
    doc_lines.append("---")
    doc_lines.append("")
    
    # 按目录结构递归添加内容
    if outline_data and 'outline' in outline_data:
        for chapter in outline_data['outline']:
            _add_chapter_content_recursive(chapter, organized_content, doc_lines, level=1)
    
    return "\n".join(doc_lines)


def _add_chapter_content_recursive(chapter, organized_content, doc_lines, level=1):
    """
    递归添加章节内容到文档
    
    Args:
        chapter: 当前章节
        organized_content: 组织后的内容
        doc_lines: 文档行列表
        level: 当前层级
    """
    chapter_id = chapter.get('id', '')
    chapter_title = chapter.get('title', '未命名章节')
    chapter_description = chapter.get('description', '')
    
    # 添加章节标题
    title_prefix = "#" * min(level + 1, 6) + " "
    doc_lines.append(f"{title_prefix}{chapter_id} {chapter_title}")
    doc_lines.append("")
    
    # 检查是否为叶子节点
    is_leaf = not chapter.get('children') or len(chapter['children']) == 0
    
    if is_leaf:
        # 叶子节点，添加生成的内容
        if chapter_id in organized_content:
            content = organized_content[chapter_id]['content']
            if content and not content.startswith('生成'):  # 排除错误信息
                doc_lines.append(content)
                doc_lines.append("")
        else:
            # 如果没有生成内容，添加描述
            if chapter_description:
                doc_lines.append(f"*{chapter_description}*")
                doc_lines.append("")
    else:
        # 非叶子节点，添加描述（如果有）并递归处理子节点
        if chapter_description:
            doc_lines.append(f"*{chapter_description}*")
            doc_lines.append("")
        
        # 递归处理子节点
        for child in chapter['children']:
            _add_chapter_content_recursive(child, organized_content, doc_lines, level + 1)


def _process_outline_recursively(data, openai_service, project_overview, progress_bar, 
                                status_text, total_nodes, processed_count, current_path):
    """
    递归处理outline_data，为叶子节点生成内容
    
    Args:
        data: 当前处理的数据节点
        openai_service: OpenAI服务实例
        project_overview: 项目概述
        progress_bar: 进度条组件
        status_text: 状态文本组件
        total_nodes: 总叶子节点数
        processed_count: 已处理计数（字典形式，用于递归修改）
        current_path: 当前路径
    """
    # 处理outline根节点
    for i, chapter in enumerate(data['outline']):
        chapter_path = f"{chapter['id']} {chapter['title']}"
        _process_single_node(chapter, openai_service, project_overview, progress_bar, 
                            status_text, total_nodes, processed_count, chapter_path, data, 
                            level=1, parent_chapters=[], all_siblings=data['outline'])


def _process_single_node(chapter, openai_service, project_overview, progress_bar, 
                        status_text, total_nodes, processed_count, chapter_path, outline_data, level=3, parent_chapters=None, all_siblings=None):
    """
    处理单个节点，判断是否为叶子节点并生成内容
    
    Args:
        level: 标题级别，用于生成对应的markdown标题（1-6级）
    """
    # 根据层级生成相应的标题
    
    # title_prefix = "#" * min(level, 6) + " "
    # st.markdown(title_prefix + chapter['title'])
    st.markdown("#" * (level + 2)+" " + chapter['title'])
    # 添加调试信息（可选）
    # st.write(f"DEBUG: 处理章节 {chapter['id']} - {chapter['title']}, 层级: {level}")
    
    # 检查是否为叶子节点（没有children或children为空）
    is_leaf = not chapter.get('children') or len(chapter['children']) == 0
    
    if is_leaf:
        # 叶子节点，生成内容
        processed_count["count"] += 1
        current_progress = processed_count["count"] / total_nodes
        progress_bar.progress(current_progress)
        status_text.markdown(f"**正在生成第 {processed_count['count']}/{total_nodes} 个章节:**  \n{chapter_path}")
        
        # 直接显示正文内容，不使用expander
        generated_text = ""
        try:
            # 创建一个文本容器用于实时显示
            #text_container = st.empty()
            
            # 获取兄弟章节信息（同级其他章节）
            sibling_chapters = []
            if all_siblings:
                sibling_chapters = [sib for sib in all_siblings if sib.get('id') != chapter.get('id')]
            
            # 直接为当前叶子节点生成内容
            generated_text = openai_service._generate_chapter_content(
                chapter, 
                parent_chapters=parent_chapters if parent_chapters else [], 
                sibling_chapters=sibling_chapters, 
                project_overview=project_overview
            )
            st.markdown(generated_text)
            
            # 保存生成的内容
            st.session_state.generated_content[chapter_path] = generated_text
            
        except Exception as e:
            error_msg = f"生成失败: {str(e)}"
            st.error(error_msg)
            st.session_state.generated_content[chapter_path] = error_msg
        
        # 确保继续处理下一个章节（无论当前章节是否成功）
    else:
        # 非叶子节点，递归处理子节点
        # 构建当前章节信息，添加到parent_chapters中
        current_parent_info = {
            'id': chapter.get('id'),
            'title': chapter.get('title'), 
            'description': chapter.get('description', '')
        }
        new_parent_chapters = (parent_chapters or []) + [current_parent_info]
        
        for child in chapter['children']:
            child_path = f"{chapter_path}/{child['id']} {child['title']}"
            _process_single_node(child, openai_service, project_overview, progress_bar, 
                               status_text, total_nodes, processed_count, child_path, outline_data, 
                               level=level+1, parent_chapters=new_parent_chapters, all_siblings=chapter['children'])