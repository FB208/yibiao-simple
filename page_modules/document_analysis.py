import streamlit as st
from typing import Dict, Optional, Tuple
import PyPDF2
import docx
import io
import logging
import queue
import time
from concurrent.futures import ThreadPoolExecutor
from enum import Enum

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class FileType(Enum):
    PDF = "application/pdf"
    DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

class SessionManager:
    """集中管理Streamlit的session state"""
    
    KEYS = {
        'api_key': None,
        'file_content': '',
        'start_analysis': False,
        'project_overview': '',
        'tech_requirements': ''
    }
    
    @staticmethod
    def initialize():
        """初始化所有session state"""
        for key, default_value in SessionManager.KEYS.items():
            if key not in st.session_state:
                st.session_state[key] = default_value
                
    @staticmethod
    def get(key: str):
        return st.session_state.get(key)
    
    @staticmethod
    def set(key: str, value):
        st.session_state[key] = value
        
    @staticmethod
    def update(updates: Dict):
        for key, value in updates.items():
            SessionManager.set(key, value)

class MarkdownEditor:
    """可编辑的Markdown区域组件"""
    
    def __init__(self, title: str, session_key: str, placeholder: str):
        self.title = title
        self.session_key = session_key
        self.placeholder = placeholder
        self.edit_mode_key = f"{session_key}_edit_mode"
        
        # 初始化session state
        if self.edit_mode_key not in st.session_state:
            st.session_state[self.edit_mode_key] = False
        if self.session_key not in st.session_state:
            st.session_state[self.session_key] = ""
            
    def render(self):
        st.markdown(f"**{self.title}**")
        self._render_mode_switch()
        
        if st.session_state[self.edit_mode_key]:
            self._render_editor()
        else:
            self._render_preview()

            
    def _render_mode_switch(self):
        col_edit, col_preview = st.columns(2)
        edit_mode = st.session_state[self.edit_mode_key]
        
        with col_edit:
            if edit_mode:
                st.button("✏️ 编辑中...", key=f"{self.session_key}_edit_btn", use_container_width=True, disabled=True, type="primary")
            else:
                if st.button("✏️ 编辑", key=f"{self.session_key}_edit_btn", use_container_width=True):
                    st.session_state[self.edit_mode_key] = True
                    st.rerun()
        
        with col_preview:
            if not edit_mode:
                st.button("👁️ 预览中...", key=f"{self.session_key}_preview_btn", use_container_width=True, disabled=True, type="primary")
            else:
                if st.button("👁️ 预览", key=f"{self.session_key}_preview_btn", use_container_width=True):
                    st.session_state[self.edit_mode_key] = False
                    st.rerun()
                    
    def _render_editor(self):
        new_value = st.text_area(
            f"{self.title} (支持Markdown格式)",
            value=st.session_state.get(self.session_key, ''),
            height=300,
            key=f"{self.session_key}_input",
            placeholder=self.placeholder
        )
        if new_value != st.session_state[self.session_key]:
            st.session_state[self.session_key] = new_value
            
    def _render_preview(self):
        content = st.session_state.get(self.session_key)
        if content:
            st.markdown(f"""
            <div style=\"border: 1px solid #e6e6e6; border-radius: 5px; padding: 15px; margin-bottom: 20px;\">
                <h4 style=\"margin-bottom: 10px;\">{self.title} 预览</h4>
                <hr>
                {content}
            </div>
            """, unsafe_allow_html=True)




class DocumentProcessor:
    """文档处理类，负责各种格式文档的文本提取"""
    
    @staticmethod
    def extract_text_from_file(uploaded_file) -> Tuple[bool, str]:
        """
        从上传的文件中提取文本内容
        
        Args:
            uploaded_file: 上传的文件对象
            
        Returns:
            Tuple[bool, str]: (是否成功, 文本内容或错误信息)
        """
        if uploaded_file is None:
            return False, "未上传文件"
            
        try:
            # 获取文件内容（避免多次读取）
            file_content = uploaded_file.getvalue()
            
            if uploaded_file.type == FileType.PDF.value:
                return DocumentProcessor._extract_pdf_text(file_content)
            elif uploaded_file.type == FileType.DOCX.value:
                return DocumentProcessor._extract_docx_text(file_content)
            else:
                supported_types = [FileType.PDF.value, FileType.DOCX.value]
                return False, f"不支持的文件格式。支持的格式：{', '.join(supported_types)}"
                
        except Exception as e:
            logging.error(f"文件解析错误: {str(e)}")
            return False, f"文件解析错误: {str(e)}"
    
    @staticmethod
    def _extract_pdf_text(file_content: bytes) -> Tuple[bool, str]:
        """提取PDF文本"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            if len(pdf_reader.pages) == 0:
                return False, "PDF文件为空或无法读取"
                
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text().strip()
                    if page_text:  # 只添加非空页面
                        text_parts.append(f"\n--- 第{page_num}页 ---\n{page_text}")
                except Exception as e:
                    logging.warning(f"无法提取第{page_num}页内容: {str(e)}")
                    continue
                    
            if not text_parts:
                return False, "PDF文件中未找到可提取的文本内容"
                
            return True, "\n".join(text_parts)
            
        except Exception as e:
            return False, f"PDF解析失败: {str(e)}"
    
    @staticmethod
    def _extract_docx_text(file_content: bytes) -> Tuple[bool, str]:
        """提取Word文档文本"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 只添加非空段落
                    text_parts.append(text)
                    
            # 处理表格内容
            for table in doc.tables:
                table_text = DocumentProcessor._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n--- 表格内容 ---\n{table_text}")
                    
            if not text_parts:
                return False, "Word文档中未找到可提取的文本内容"
                
            return True, "\n\n".join(text_parts)
            
        except Exception as e:
            return False, f"Word文档解析失败: {str(e)}"
    
    @staticmethod
    def _extract_table_text(table) -> str:
        """提取表格文本"""
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text if cell_text else "-")
                if any(cell != "-" for cell in row_data):  # 只添加非空行
                    table_data.append(" | ".join(row_data))
            return "\n".join(table_data)
        except Exception:
            return ""



def render_document_analysis_page() -> Dict:
    """渲染标书解析页面"""
    st.markdown("### 📄 第一步：上传并解析招标文件")
    _render_markdown_guide()
    
    # 初始化session
    SessionManager.initialize()
    
    # 文件上传
    uploaded_file = _render_file_uploader()
    
    if uploaded_file:
        # 检查是否正在解析中，如果是则禁用按钮
        is_analyzing = SessionManager.get('start_analysis')
        button_text = "⏳ 解析中..." if is_analyzing else "🔍 开始解析"
        if st.button(button_text, type="primary", use_container_width=True, disabled=is_analyzing):
            _handle_analysis_request(uploaded_file)
    
    # 解析结果预览
    st.markdown("---")
    st.subheader("📋 解析结果预览")
    _render_analysis_results()
    
    # 内容总结
    _render_summary()
    
    return {
        'project_overview': SessionManager.get('project_overview'),
        'tech_requirements': SessionManager.get('tech_requirements'),
        'uploaded_file': uploaded_file
    }

def _render_markdown_guide():
    """渲染Markdown编辑说明"""
    with st.expander("📝 Markdown编辑说明", expanded=False):
        st.markdown("""
        **支持以下Markdown格式：**
        
        - **粗体文本**：`**文本**`
        - *斜体文本*：`*文本*`
        - `代码`：`` `代码` ``
        - [链接](URL)：`[链接文本](URL)`
        
        **列表和引用：**
        - 无序列表：`- 项目`
        - 有序列表：`1. 项目`
        > 引用文本：`> 文本`
        """)

def _render_file_uploader():
    """渲染文件上传区域"""
    return st.file_uploader(
        "上传招标文件",
        type=['docx', 'pdf'],
        help="支持Word (.docx) 和 PDF (.pdf) 格式",
        key="doc_analysis_file"
    )

def _handle_analysis_request(uploaded_file):
    """处理文档解析请求"""
    if not SessionManager.get('api_key'):
        st.error("请先在左侧配置面板中设置OpenAI API密钥！")
        return

    with st.spinner("正在解析文档..."):
        success, content = DocumentProcessor.extract_text_from_file(uploaded_file)
        
        if success:
            SessionManager.set('file_content', content)
            st.success("文档解析完成！正在生成分析结果...")
            SessionManager.set('start_analysis', True)
            st.rerun()
        else:
            st.error(f"文档解析失败: {content}")

def _fetch_analysis_stream(analysis_type: str, result_queue: queue.Queue, openai_service, file_content: str):
    """在线程中获取AI分析结果并放入队列"""
    try:
        content_stream = openai_service.analyze_document(file_content, analysis_type)
        for chunk in content_stream:
            result_queue.put((analysis_type, chunk))
    except Exception as e:
        result_queue.put((analysis_type, f"AI分析失败: {str(e)}"))
    finally:
        result_queue.put((analysis_type, None))

def _render_analysis_results():
    """渲染AI分析结果"""

    col1, col2 = st.columns(2)

    overview_editor = MarkdownEditor(
        title="项目概述",
        session_key="project_overview",
        placeholder="这里将显示项目概述内容..."
    )
    requirements_editor = MarkdownEditor(
        title="技术评分要求",
        session_key="tech_requirements",
        placeholder="这里将显示技术评分要求内容..."
    )

    if not SessionManager.get('start_analysis'):
        with col1:
            overview_editor.render()
        with col2:
            requirements_editor.render()
        return

    # --- Concurrent Analysis Logic ---
    try:
        from services.openai_servce import get_openai_service
        openai_service = get_openai_service()
        file_content = SessionManager.get('file_content')
    except Exception as e:
        st.error(f"初始化服务失败: {e}")
        SessionManager.set('start_analysis', False)
        st.rerun()
        return
        
    result_queue = queue.Queue()
    analysis_map = {
        "overview": {"editor": overview_editor, "content": "", "finished": False},
        "requirements": {"editor": requirements_editor, "content": "", "finished": False},
    }

    # Render placeholders
    with col1:
        st.markdown(f"**{overview_editor.title}**")
        overview_editor._render_mode_switch()
        overview_placeholder = st.empty()
    with col2:
        st.markdown(f"**{requirements_editor.title}**")
        requirements_editor._render_mode_switch()
        req_placeholder = st.empty()
    
    placeholders = {"overview": overview_placeholder, "requirements": req_placeholder}

    with ThreadPoolExecutor(max_workers=2) as executor:
        executor.submit(_fetch_analysis_stream, "overview", result_queue, openai_service, file_content)
        executor.submit(_fetch_analysis_stream, "requirements", result_queue, openai_service, file_content)

        while not all(d["finished"] for d in analysis_map.values()):
            try:
                analysis_type, chunk = result_queue.get(timeout=30)
                data = analysis_map[analysis_type]
                
                if chunk is None:
                    data["finished"] = True
                    continue

                if isinstance(chunk, str):
                    data["content"] += chunk
                
                content_html = f'''
                <div style="border: 1px solid #e6e6e6; border-radius: 5px; padding: 15px; margin-bottom: 20px;">
                    <h4 style="margin-bottom: 10px;">{data["editor"].title} 预览</h4>
                    <hr>
                    {data["content"]}
                </div>
                '''
                placeholders[analysis_type].markdown(content_html, unsafe_allow_html=True)
            except queue.Empty:
                st.warning("分析超时，请检查网络或API设置。")
                SessionManager.set('start_analysis', False)
                break
            except Exception as e:
                st.error(f"处理队列时出错: {e}")
                SessionManager.set('start_analysis', False)
                break
    
    SessionManager.update({
        'project_overview': analysis_map['overview']['content'],
        'tech_requirements': analysis_map['requirements']['content'],
        'start_analysis': False
    })
    
    st.success("分析全部完成！")
    time.sleep(1)
    st.rerun()

def _render_summary():
    """渲染内容总结区域"""
    overview = SessionManager.get('project_overview')
    requirements = SessionManager.get('tech_requirements')
    
    if overview or requirements:
        st.markdown("---")
        st.subheader("📊 当前内容总结")
        
        col1, col2 = st.columns(2)
        if overview:
            with col1:
                st.markdown(f"**项目概述字数：** {len(overview)}")
        if requirements:
            with col2:
                st.markdown(f"**技术要求字数：** {len(requirements)}")
